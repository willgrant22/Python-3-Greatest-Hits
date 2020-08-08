#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# =============================================================================
# Author :  Will Grant
# =============================================================================

import docx 

doc = docx.Document() 

doc.add_heading('Heading for the document', 0) 

doc_para = doc.add_paragraph('Your paragraph goes here, ') 

doc_para.add_run('hey there, bold here').bold = True
doc_para.add_run(', and ') 
doc_para.add_run('these words are italic').italic = True

doc.add_page_break() 

doc.add_heading('Heading level 2', 2) 

doc.save('test.docx') 

